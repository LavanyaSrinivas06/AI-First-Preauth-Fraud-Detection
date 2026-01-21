#!/usr/bin/env python3
"""Convert a simple Markdown file to DOCX (images & headings supported).

Usage: python scripts/md_to_docx.py docs/phase_id_full_report.md docs/phase_id_full_report.docx

This script implements a small subset of Markdown: headings, paragraphs, and image embedding.
It resolves image paths relative to the markdown file directory.
"""
import sys
from pathlib import Path
import re
from docx import Document
from docx.shared import Inches

IMG_RE = re.compile(r'!\[(.*?)\]\((.*?)\)')
HEADING_RE = re.compile(r'^(#{1,6})\s+(.*)')


def md_to_docx(md_path: Path, out_path: Path):
    md_path = Path(md_path)
    out_path = Path(out_path)
    if not md_path.exists():
        raise SystemExit(f"Markdown file not found: {md_path}")

    doc = Document()
    cur_para = None

    with md_path.open(encoding="utf-8") as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')
        if not line.strip():
            # blank line - add spacing
            cur_para = None
            i += 1
            continue

        # Heading
        m = HEADING_RE.match(line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            style = None
            if level == 1:
                style = 'Heading 1'
            elif level == 2:
                style = 'Heading 2'
            elif level == 3:
                style = 'Heading 3'
            p = doc.add_paragraph(text, style=style if style else None)
            cur_para = p
            i += 1
            continue

        # Image
        m = IMG_RE.search(line)
        if m:
            alt = m.group(1) or ''
            img_rel = m.group(2)
            # resolve relative to the markdown file
            img_path = (md_path.parent / img_rel).resolve()
            if not img_path.exists():
                # try absolute as given
                img_path = Path(img_rel)
            if img_path.exists():
                # add caption (alt text) then image
                if alt:
                    doc.add_paragraph(alt, style='Intense Quote')
                try:
                    doc.add_picture(str(img_path), width=Inches(6))
                except Exception:
                    # fallback: smaller width
                    try:
                        doc.add_picture(str(img_path), width=Inches(4))
                    except Exception:
                        doc.add_paragraph(f"[Image not embedded: {img_path}]")
            else:
                doc.add_paragraph(f"[Missing image: {img_rel}]")
            cur_para = None
            i += 1
            continue

        # Fenced code block start
        if line.strip().startswith('```'):
            code_lines = []
            fence = line.strip()
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i].rstrip('\n'))
                i += 1
            # skip closing fence
            i += 1
            # add as preformatted paragraph
            p = doc.add_paragraph('\n'.join(code_lines))
            p.style = 'Intense Quote'
            cur_para = None
            continue

        # Normal paragraph or list
        # accumulate subsequent lines that are not blank, heading, or image
        para_lines = [line]
        j = i + 1
        while j < len(lines):
            nxt = lines[j].rstrip('\n')
            if not nxt.strip():
                break
            if HEADING_RE.match(nxt) or IMG_RE.search(nxt) or nxt.strip().startswith('```'):
                break
            para_lines.append(nxt)
            j += 1
        text = ' '.join(l.strip() for l in para_lines)
        doc.add_paragraph(text)
        cur_para = None
        i = j

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"Wrote {out_path}")


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python scripts/md_to_docx.py input.md output.docx")
        raise SystemExit(2)
    md_to_docx(Path(sys.argv[1]), Path(sys.argv[2]))
